"""
ContextOS ingestors/docx.py — Word document extraction via python-docx.

Converts DOCX structure to Markdown:
  Heading 1 → # heading
  Heading 2 → ## heading
  Heading 3 → ### heading
  Normal paragraph → plain text
  Table → Markdown table
  Bold/Italic inline styles preserved via ** and *
"""
from __future__ import annotations

import time
from pathlib import Path

from contextos.ingestors import register


def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    lines = []
    header = [_cell_text(c) for c in rows[0].cells]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(["---"] * len(header)) + "|")
    for row in rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _para_to_md(para) -> str:
    """Convert a paragraph to Markdown, preserving heading level and basic inline styles."""
    style = para.style.name if para.style else ""
    text  = para.text.strip()

    if not text:
        return ""

    # Heading styles
    if style.startswith("Heading 1"): return f"# {text}"
    if style.startswith("Heading 2"): return f"## {text}"
    if style.startswith("Heading 3"): return f"### {text}"
    if style.startswith("Heading 4"): return f"#### {text}"
    if style.startswith("Title"):     return f"# {text}"
    if style.startswith("Subtitle"):  return f"*{text}*"

    # List styles
    if "List Bullet" in style: return f"- {text}"
    if "List Number" in style: return f"1. {text}"

    # Inline bold/italic from runs
    parts = []
    for run in para.runs:
        rt = run.text
        if not rt:
            continue
        if run.bold and run.italic:
            parts.append(f"***{rt}***")
        elif run.bold:
            parts.append(f"**{rt}**")
        elif run.italic:
            parts.append(f"*{rt}*")
        else:
            parts.append(rt)
    return "".join(parts) if parts else text


@register(".docx", "python-docx")
def extract_docx(path: Path) -> str:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))

    # Try to get title from core properties
    try:
        cp    = doc.core_properties
        title = (cp.title or "").strip() or path.stem.replace("-"," ").replace("_"," ").title()
        author = (cp.author or "").strip() or None
    except Exception:
        title  = path.stem.replace("-"," ").replace("_"," ").title()
        author = None

    fm_lines = [
        "---",
        "project: unknown",
        "type: note",
        "status: draft",
        "source: docx",
        f"updated_at: {time.strftime('%Y-%m-%d')}",
        "tags:",
        "  - docx",
        "  - imported",
    ]
    if author:
        fm_lines.append(f"owner: {author}")
    fm_lines += ["---", "", f"# {title}", ""]

    # Iterate body elements preserving order of paragraphs and tables
    body    = doc.element.body
    content = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Find matching paragraph object
            for para in doc.paragraphs:
                if para._element is child:
                    md = _para_to_md(para)
                    if md:
                        content.append(md)
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is child:
                    content.append(_table_to_md(table))
                    break

    return "\n".join(fm_lines) + "\n" + "\n\n".join(content)
